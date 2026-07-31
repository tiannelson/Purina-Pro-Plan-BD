import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TITLE_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
HEADING_COLOR = RGBColor(0x2E, 0x5E, 0x8C)

def set_run(run, size=11, bold=False, italic=False, color=None, font="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    if color:
        run.font.color.rgb = color

def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=22, bold=True, color=TITLE_COLOR)
    # bottom border under title block
    return p

def add_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pPr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '8', qn('w:space'): '1', qn('w:color'): '2E5E8C'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_byline(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    set_run(run, size=10, italic=True, color=RGBColor(0x44, 0x44, 0x44))

def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, size=13, bold=True, color=HEADING_COLOR)

def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run(r1, size=10.5, bold=True)
            r2 = p.add_run(rest)
            set_run(r2, size=10.5)
        else:
            r = p.add_run(item)
            set_run(r, size=10.5)

def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(6)
        if isinstance(item, tuple):
            lead, rest = item
            r1 = p.add_run(lead)
            set_run(r1, size=10.5, bold=True)
            r2 = p.add_run(rest)
            set_run(r2, size=10.5)
        else:
            r = p.add_run(item)
            set_run(r, size=10.5)

def build_memo(out_path, title, byline, recap, signals, implications, actions):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc, title)
    add_rule(doc)
    add_byline(doc, byline)

    add_heading(doc, "1. Meeting Recap")
    add_bullets(doc, recap)

    add_heading(doc, "2. Market Signals")
    add_bullets(doc, signals)

    add_heading(doc, "3. Implications")
    add_bullets(doc, implications)

    add_heading(doc, "4. Recommended Actions")
    add_numbered(doc, actions)

    doc.save(out_path)
    print("wrote", out_path)


with open(os.path.join(os.path.dirname(__file__), "memos.json")) as f:
    memos = json.load(f)

outdir = os.path.dirname(__file__)
for m in memos:
    recap = [tuple(x) if isinstance(x, list) else x for x in m["recap"]]
    signals = [tuple(x) if isinstance(x, list) else x for x in m["signals"]]
    implications = [tuple(x) if isinstance(x, list) else x for x in m["implications"]]
    actions = [tuple(x) if isinstance(x, list) else x for x in m["actions"]]
    build_memo(
        os.path.join(outdir, m["filename"]),
        m["title"],
        m["byline"],
        recap, signals, implications, actions,
    )
